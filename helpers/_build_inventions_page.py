#!/usr/bin/env python3
"""Build live Major Scientific Inventions HTML from the inventions DOCX."""
from __future__ import annotations

import html
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from _paths import ROOT
from _kt_favicon import favicon_links
from _inventions_catalog_filters import (
    infer_period,
    render_category_select,
    render_period_select,
)

DOCX = ROOT / "documents" / "Most_Influential_Scientific_Inventions_and_Innovations (with references).docx"
OUT_HTML = ROOT / "en" / "major_scientific_inventions.html"
OUT_HTML_AZ = ROOT / "az" / "major_scientific_inventions.html"
OUT_IMAGES = ROOT / "images" / "inventions"
CARD_DATA = ROOT / "preview" / "inventions-card-data.json"
CARD_OVERRIDES = ROOT / "preview" / "inventions-card-overrides.json"
CARD_COPY = ROOT / "helpers" / "_inventions_card_copy.json"
LINK_FIXES = ROOT / "helpers" / "_invention_link_fixes.json"
ENTRY_REFERENCES = ROOT / "helpers" / "_invention_entry_references.json"
ICON_INDEX = ROOT / "helpers" / "_invention_icon_index.json"
NAV_SOURCE = ROOT / "en" / "index.html"
ILLUSTRATION_SECTION = "6"
INVENTIONS_DIR = ROOT / "images" / "inventions"
ICONS_DIR = ROOT / "images" / "icons"
IMAGE_ALIASES: dict[str, str] = {
    "integrated-circuit-microchip": "microchip",
    "fibre-optic-communication": "fibre-optic",
}
ICON_FILENAME_CANDIDATES: dict[str, list[str]] = {
    "the-wheel": ["the-wheel.png", "wheel.png"],
    "the-compass": ["the-compass.png", "compass.png"],
    "the-printing-press": ["the-printing-press.png", "printing-press.png"],
    "the-scientific-method": ["the-scientific-method.png", "scientific-method.png"],
    "the-telescope": ["the-telescope.png", "telescope.png"],
    "the-microscope": ["the-microscope.png", "microscope.png"],
    "the-periodic-table": ["the-periodic-table.png", "periodic-table.png"],
    "integrated-circuit-microchip": ["integrated-circuit-microchip.png", "microchip.png"],
    "fibre-optic-communication": ["fibre-optic-communication.png", "fibre-optic.png"],
}

SECTION_LABELS = {
    "HISTORICAL CONTEXT": "Historical context",
    "SCIENTIFIC / TECHNOLOGICAL SIGNIFICANCE": "Scientific significance",
    "IMPACT ON HUMANITY": "Impact on humanity",
    "MULTIMEDIA & FURTHER RESOURCES": "Further resources",
}

SKIP_H1 = {"contents", "overview by category"}

URL_RE = re.compile(r"https?://[^\s\]\)>,]+")


def slugify(text: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return s or "entry"


def esc(text: str) -> str:
    return html.escape(text, quote=True)


def load_link_fixes() -> tuple[dict[str, str], set[str]]:
    if not LINK_FIXES.exists():
        return {}, set()
    payload = json.loads(LINK_FIXES.read_text(encoding="utf-8"))
    replacements: dict[str, str] = {}
    removals = set(payload.get("removals", []))
    for old, new in payload.get("replacements", {}).items():
        if new:
            replacements[old] = new
        else:
            removals.add(old)
    return replacements, removals


def fix_url(url: str) -> str | None:
    if not url:
        return url
    url = url.rstrip(".,;)]")
    replacements, removals = load_link_fixes()
    if url in removals:
        return None
    return replacements.get(url, url)


def apply_text_link_fixes(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        raw = match.group(0)
        url = raw.rstrip(".,;)]")
        suffix = raw[len(url) :]
        fixed = fix_url(url)
        if fixed is None:
            return ""
        return fixed + suffix

    cleaned = URL_RE.sub(repl, text)
    return re.sub(r"  +", " ", cleaned).strip()


def apply_link_fixes(data: dict) -> None:
    for entry in data["entries"]:
        for section_name, items in entry.get("sections", {}).items():
            if section_name != "MULTIMEDIA & FURTHER RESOURCES":
                continue
            kept: list[dict[str, str] | str] = []
            for item in items:
                if not isinstance(item, dict):
                    kept.append(apply_text_link_fixes(item) if isinstance(item, str) else item)
                    continue
                url = item.get("url", "")
                if url:
                    fixed = fix_url(url)
                    if fixed is None:
                        continue
                    item = {**item, "url": fixed}
                kept.append(item)
            entry["sections"][section_name] = kept

    refs = data.get("references", {})
    for group in refs.get("groups", []):
        group["items"] = [apply_text_link_fixes(item) for item in group.get("items", [])]


def linkify_text(text: str) -> str:
    parts: list[str] = []
    last = 0
    for match in URL_RE.finditer(text):
        parts.append(esc(text[last : match.start()]))
        url = match.group(0).rstrip(".,;)]")
        suffix = match.group(0)[len(url) :]
        fixed = fix_url(url)
        if fixed is not None:
            parts.append(
                f'<a class="resource-link" href="{esc(fixed)}" target="_blank" '
                f'rel="noopener noreferrer">{esc(fixed)}</a>{esc(suffix)}'
            )
        last = match.end()
    parts.append(esc(text[last:]))
    return "".join(parts)


def parse_resource_line(line: str) -> tuple[str, str]:
    line = line.lstrip("▸ ").strip()
    if " — " in line:
        parts = [p.strip() for p in line.split(" — ")]
        if len(parts) >= 2:
            return parts[0], parts[-1]
    return line, ""


def extract_para_url(para: Paragraph) -> str:
    for hyperlink in para._element.iter(qn("w:hyperlink")):
        rel_id = hyperlink.get(qn("r:id"))
        if rel_id and para.part.rels.get(rel_id):
            return para.part.rels[rel_id].target_ref
    return ""


def extract_para_image(para: Paragraph) -> str:
    for blip in para._element.findall(".//" + qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if rel_id and para.part.rels.get(rel_id):
            return para.part.rels[rel_id].target_ref
    return ""


def parse_resource_item(para: Paragraph) -> dict[str, str]:
    text = para.text.strip()
    title, source = parse_resource_line(text)
    return {"title": title, "source": source, "url": extract_para_url(para)}


def parse_docx() -> dict:
    doc = Document(DOCX)
    data: dict = {
        "title": "",
        "subtitle": "",
        "tagline": "",
        "executive_summary": "",
        "categories": [],
        "entries": [],
        "overview_table": [],
        "top20_table": [],
        "conclusion": {"title": "", "paragraphs": []},
        "references": {"title": "", "intro": "", "groups": []},
    }

    paras = doc.paragraphs
    if paras:
        data["title"] = paras[0].text.strip()
    if len(paras) > 1:
        data["subtitle"] = paras[1].text.strip()
    if len(paras) > 2:
        data["tagline"] = paras[2].text.strip()
    for para in paras[3:8]:
        t = para.text.strip()
        if t.upper().startswith("EXECUTIVE SUMMARY"):
            data["executive_summary"] = re.sub(
                r"^EXECUTIVE SUMMARY\s*",
                "",
                t,
                flags=re.IGNORECASE,
            ).strip()
            break

    current_cat: dict | None = None
    current_entry: dict | None = None
    current_section: str | None = None
    current_ref_group: dict | None = None
    in_conclusion = False
    in_references = False

    def flush_entry() -> None:
        nonlocal current_entry
        if current_entry and current_cat:
            current_cat["entries"].append(current_entry)
        current_entry = None

    for para in paras:
        text = para.text.strip()
        style = para.style.name if para.style else "Normal"

        if (
            current_entry is not None
            and not in_conclusion
            and not in_references
            and style != "Heading 2"
        ):
            image_ref = extract_para_image(para)
            if image_ref:
                current_entry["image"] = image_ref
                continue

        if not text:
            continue

        if style == "Heading 1":
            key = text.lower().lstrip("0123456789. ").strip()
            if key in SKIP_H1:
                current_entry = None
                current_section = None
                in_conclusion = False
                continue
            if key.startswith("concise top 20"):
                current_cat = None
                current_entry = None
                in_conclusion = False
                continue
            if key == "concluding note":
                in_conclusion = True
                in_references = False
                current_cat = None
                current_entry = None
                current_ref_group = None
                data["conclusion"]["title"] = text
                continue
            if key == "references":
                in_references = True
                in_conclusion = False
                current_cat = None
                current_entry = None
                current_ref_group = None
                data["references"]["title"] = text
                continue
            in_conclusion = False
            in_references = False
            current_ref_group = None
            flush_entry()
            text = normalize_category_title(text)
            current_cat = {"title": text, "slug": slugify(text), "entries": []}
            data["categories"].append(current_cat)
            current_section = None
            continue

        if in_conclusion:
            data["conclusion"]["paragraphs"].append(text)
            continue

        if in_references:
            if style == "Heading 3":
                text = normalize_ref_group_title(text)
                current_ref_group = {
                    "title": text,
                    "slug": slugify(text),
                    "items": [],
                }
                data["references"]["groups"].append(current_ref_group)
                continue
            if style == "Normal":
                if current_ref_group is None and not data["references"]["intro"]:
                    data["references"]["intro"] = text
                elif current_ref_group is not None:
                    current_ref_group["items"].append(text)
            continue

        if style == "Heading 2":
            flush_entry()
            current_entry = {
                "title": text,
                "slug": slugify(text),
                "meta": "",
                "image": "",
                "image_src": "",
                "sections": {},
                "category": current_cat["title"] if current_cat else "",
            }
            data["entries"].append(current_entry)
            current_section = None
            continue

        if current_entry is None:
            continue

        if text in SECTION_LABELS:
            current_section = text
            current_entry["sections"].setdefault(current_section, [])
            continue

        if current_section:
            if current_section == "MULTIMEDIA & FURTHER RESOURCES" and text.startswith("▸"):
                current_entry["sections"][current_section].append(parse_resource_item(para))
            else:
                current_entry["sections"][current_section].append(text)
        elif not current_entry["meta"] and text.lower().startswith("period:"):
            current_entry["meta"] = normalize_meta_line(text)

    flush_entry()

    for cat_index, cat in enumerate(data["categories"], start=1):
        cat_num = _cat_number(cat["title"]) or str(cat_index)
        for entry_index, entry in enumerate(cat["entries"], start=1):
            entry["number"] = f"{cat_num}.{entry_index}"

    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not rows:
            continue
        header = [h.lower() for h in rows[0]]
        if "category" in header[0] and "items" in " ".join(header):
            data["overview_table"] = rows
        elif "rank" in header[0] and "breakthrough" in " ".join(header):
            data["top20_table"] = rows

    return data


def attach_orphan_entries(data: dict, cards: dict[str, dict]) -> None:
    """Place parsed entries that missed category assignment (e.g. after References in DOCX)."""
    placed = {entry["slug"] for cat in data["categories"] for entry in cat["entries"]}
    orphans = [entry for entry in data["entries"] if entry["slug"] not in placed]
    if not orphans:
        return

    cat_by_num: dict[str, dict] = {}
    for cat in data["categories"]:
        num = _cat_number(cat["title"])
        if num:
            cat_by_num[num] = cat

    for entry in orphans:
        section_num = ""
        icon = cards.get(entry["slug"], {}).get("icon", "")
        m = re.search(r"icons/(\d+)-", icon.replace("\\", "/"))
        if m:
            section_num = m.group(1)
        if not section_num:
            for cat in data["categories"]:
                if entry.get("category") and entry["category"] == cat["title"]:
                    section_num = _cat_number(cat["title"])
                    break
        target = cat_by_num.get(section_num) or data["categories"][-1]
        target["entries"].append(entry)

    for cat_index, cat in enumerate(data["categories"], start=1):
        cat_num = _cat_number(cat["title"]) or str(cat_index)
        for entry_index, entry in enumerate(cat["entries"], start=1):
            entry["number"] = f"{cat_num}.{entry_index}"


def meta_key_figures(meta: str) -> str:
    if "|" not in meta:
        return ""
    tail = meta.split("|", 1)[1].strip()
    return re.sub(r"^Key figure\(s\):\s*", "", tail, flags=re.IGNORECASE).strip()


def load_card_data() -> dict[str, dict]:
    if not CARD_DATA.exists():
        return {}
    return json.loads(CARD_DATA.read_text(encoding="utf-8"))


def load_card_overrides() -> dict[str, dict]:
    if not CARD_OVERRIDES.exists():
        return {}
    return json.loads(CARD_OVERRIDES.read_text(encoding="utf-8"))


def load_card_copy() -> dict[str, dict]:
    if not CARD_COPY.exists():
        return {}
    return json.loads(CARD_COPY.read_text(encoding="utf-8"))


def normalize_category_title(title: str) -> str:
    """Collapse extra whitespace after the section number (e.g. '1.  Foo' → '1. Foo')."""
    title = re.sub(r"^(\d+)\.\s+", r"\1. ", title.strip())
    return re.sub(r"  +", " ", title)


def normalize_ref_group_title(title: str) -> str:
    """Collapse extra whitespace after the letter prefix (e.g. 'A.  Books' → 'A. Books')."""
    title = re.sub(r"^([A-Z])\.\s+", r"\1. ", title.strip())
    return re.sub(r"  +", " ", title)


def normalize_meta_line(text: str) -> str:
    text = re.sub(r"\s+\|\s+", " | ", text.strip())
    return re.sub(r"  +", " ", text)


def card_key_figures_clean(text: str) -> bool:
    if not text:
        return False
    if re.search(r"KEY\s*FACT|ikebdi|CHk\s*EY|MAcasmerica|ventional|ssachusetts", text, re.IGNORECASE):
        return False
    return len(text) <= 140


def load_icon_index() -> dict[str, dict]:
    if not ICON_INDEX.exists():
        return {}
    data = json.loads(ICON_INDEX.read_text(encoding="utf-8"))
    return data.get("icons", {})


def site_icon_src(slug: str, icon_index: dict[str, dict], fallback: str = "") -> str:
    resolved = resolve_icon_site_path(slug, icon_index)
    if resolved:
        return resolved
    indexed = icon_index.get(slug, {})
    if indexed.get("site_path"):
        return indexed["site_path"]
    if fallback.startswith("images/icons/"):
        return "../" + fallback
    return fallback


def icon_filename_candidates(slug: str) -> list[str]:
    names = [f"{slug}.png"]
    for alt in ICON_FILENAME_CANDIDATES.get(slug, []):
        if alt not in names:
            names.append(alt)
    return names


def resolve_icon_site_path(slug: str, icon_index: dict[str, dict]) -> str:
    info = icon_index.get(slug, {})
    section = info.get("section")
    if section:
        folder = ICONS_DIR / section
        for name in icon_filename_candidates(slug):
            if (folder / name).is_file():
                return f"../images/icons/{section}/{name}"

    for name in icon_filename_candidates(slug):
        for match in ICONS_DIR.rglob(name):
            if match.is_file() and "orijinal" not in match.parts:
                rel = match.relative_to(ROOT).as_posix()
                return f"../{rel}"
    return ""


def attach_card_content(
    data: dict,
    cards: dict[str, dict],
    overrides: dict[str, dict] | None = None,
    card_copy: dict[str, dict] | None = None,
) -> None:
    overrides = overrides or {}
    card_copy = card_copy or {}
    icon_index = load_icon_index()
    for entry in data["entries"]:
        card = cards.get(entry["slug"], {})
        override = overrides.get(entry["slug"], {})
        clean = card_copy.get(entry["slug"], {})
        meta_figures = meta_key_figures(entry.get("meta", ""))
        ocr_figures = card.get("key_figures", "")
        entry["card"] = card
        entry["card_key_figures"] = (
            override.get("key_figures")
            or clean.get("key_figures")
            or (ocr_figures if card_key_figures_clean(ocr_figures) else meta_figures)
        )
        entry["card_summary"] = (
            override.get("summary") or clean.get("summary") or card.get("summary", "")
        )
        entry["card_facts"] = (
            override.get("key_facts") or clean.get("key_facts") or card.get("key_facts", [])
        )
        raw_icon = override.get("icon") or card.get("icon", "")
        entry["card_icon"] = site_icon_src(entry["slug"], icon_index, raw_icon)


def invention_image_src(entry: dict) -> str | None:
    slug = entry["slug"]
    names = [IMAGE_ALIASES.get(slug, slug)]
    if slug not in names:
        names.append(slug)
    for name in names:
        if (INVENTIONS_DIR / f"{name}.png").exists():
            return f"../images/inventions/{name}.png"
    return None


def render_icon_image(entry: dict, src: str) -> str:
    return (
        '<figure class="inventions-entry-icon">'
        f'<img src="{esc(src)}" alt="{esc("Illustration: " + entry["title"])}" '
        'loading="lazy" decoding="async"/>'
        f'<figcaption class="inventions-entry-icon-caption">'
        f"Illustration: {esc(entry['title'])}</figcaption>"
        "</figure>"
    )


def render_entry_icon(entry: dict) -> str:
    number = entry.get("number", "")
    if number.startswith(f"{ILLUSTRATION_SECTION}."):
        src = invention_image_src(entry)
        if src:
            return render_icon_image(entry, src)

    src = entry.get("card_icon", "")
    if src and (ROOT / src.removeprefix("../")).is_file():
        return render_icon_image(entry, src)
    return render_icon_placeholder(entry)


def render_icon_placeholder(entry: dict) -> str:
    return (
        '<figure class="inventions-entry-icon">'
        '<div class="inventions-entry-icon-placeholder" role="img" '
        f'aria-label="{esc("Illustration placeholder: " + entry["title"])}">'
        '<span class="inventions-entry-icon-placeholder__glyph" aria-hidden="true">🖼</span>'
        '<span class="inventions-entry-icon-placeholder__label">Illustration</span>'
        "</div>"
        f'<figcaption class="inventions-entry-icon-caption">'
        f"Illustration: {esc(entry['title'])}</figcaption>"
        "</figure>"
    )


def render_entry_visual(entry: dict) -> str:
    if (
        not entry.get("card_summary")
        and not entry.get("card_facts")
        and not entry.get("card_key_figures")
    ):
        return ""

    parts = ['<div class="inventions-entry-visual">', render_entry_icon(entry)]

    copy_bits: list[str] = []
    if entry.get("card_key_figures"):
        copy_bits.append(
            '<p class="inventions-entry-visual-figures">'
            f'<strong>Key figure(s):</strong> {esc(entry["card_key_figures"])}</p>'
        )
    if entry.get("card_summary"):
        copy_bits.append(
            f'<p class="inventions-entry-visual-summary">{esc(entry["card_summary"])}</p>'
        )
    if entry.get("card_facts"):
        items = "".join(f"<li>{esc(fact)}</li>" for fact in entry["card_facts"])
        copy_bits.append(
            '<div class="inventions-key-facts">'
            "<h4>Key facts</h4>"
            f"<ul>{items}</ul>"
            "</div>"
        )

    if copy_bits:
        parts.append('<div class="inventions-entry-visual-copy">' + "".join(copy_bits) + "</div>")
    parts.append("</div>")
    return "\n".join(parts)


def export_entry_images(data: dict) -> int:
    OUT_IMAGES.mkdir(parents=True, exist_ok=True)
    exported = 0
    with zipfile.ZipFile(DOCX) as archive:
        available = set(archive.namelist())
        for entry in data["entries"]:
            media_path = entry.get("image", "")
            if not media_path:
                continue
            zip_path = f"word/{media_path}"
            if zip_path not in available:
                continue
            ext = Path(media_path).suffix or ".png"
            dest = OUT_IMAGES / f"{entry['slug']}{ext}"
            dest.write_bytes(archive.read(zip_path))
            entry["image_src"] = f"images/{dest.name}"
            exported += 1
    return exported


def load_nav_html() -> str:
    text = NAV_SOURCE.read_text(encoding="utf-8")
    m = re.search(r"<nav[^>]*>.*?</nav>", text, re.DOTALL)
    return m.group(0) if m else ""


def render_resources(items: list[dict[str, str] | str]) -> str:
    lis: list[str] = []
    for item in items:
        if isinstance(item, str):
            title, source = parse_resource_line(item)
            url = ""
        else:
            title = item.get("title", "")
            source = item.get("source", "")
            url = item.get("url", "")
        if url:
            url = fix_url(url) or ""
        if url:
            label = (
                f'<a class="resource-link" href="{esc(url)}" target="_blank" '
                f'rel="noopener noreferrer">{esc(title)}</a>'
            )
        else:
            label = f'<span class="resource-title">{esc(title)}</span>'
        if source:
            lis.append(
                f"<li>{label} <span class=\"resource-source\">— {esc(source)}</span></li>"
            )
        else:
            lis.append(f"<li>{label}</li>")
    return f'<ul class="inventions-resources">{"".join(lis)}</ul>'


def load_entry_references() -> dict[str, list[str]]:
    if not ENTRY_REFERENCES.exists():
        return {}
    payload = json.loads(ENTRY_REFERENCES.read_text(encoding="utf-8"))
    return payload.get("entries", {})


def build_reference_index(data: dict) -> dict[str, dict[str, str]]:
    index: dict[str, dict[str, str]] = {}
    for group in data["references"].get("groups", []):
        for ii, item in enumerate(group.get("items", [])):
            if item.startswith("Note on Sources"):
                continue
            rid = f"{group['slug']}:{ii}"
            index[rid] = {"group": group["title"], "text": item}
    return index


def render_entry_references(
    entry: dict,
    ref_index: dict[str, dict[str, str]],
    entry_refs_map: dict[str, list[str]],
) -> str:
    ref_ids = entry_refs_map.get(entry["slug"], [])
    if not ref_ids:
        return ""

    by_group: dict[str, list[str]] = {}
    group_order: list[str] = []
    for rid in ref_ids:
        ref = ref_index.get(rid)
        if not ref:
            continue
        group = ref["group"]
        if group not in by_group:
            by_group[group] = []
            group_order.append(group)
        by_group[group].append(ref["text"])

    if not group_order:
        return ""

    parts = ['<div class="inventions-entry-references">', "<h3>References</h3>"]
    for group in group_order:
        parts.append('<div class="inventions-entry-references-group">')
        parts.append(f"<h4>{esc(group)}</h4>")
        parts.append('<ol class="inventions-bibliography">')
        for item in by_group[group]:
            parts.append(f"<li>{linkify_text(item)}</li>")
        parts.append("</ol></div>")
    parts.append("</div>")
    return "\n".join(parts)


def render_entry(
    entry: dict,
    ref_index: dict[str, dict[str, str]],
    entry_refs_map: dict[str, list[str]],
    cat_num: str = "",
) -> str:
    number = entry.get("number", "")
    card_bits = " ".join(
        [
            entry.get("card_key_figures", ""),
            entry.get("card_summary", ""),
            " ".join(entry.get("card_facts", [])),
        ]
    )
    search_blob = f"{number} {entry['title']} {entry.get('meta', '')} {card_bits}".strip()
    period = infer_period(entry.get("meta", ""), [], slug=entry.get("slug"))
    if number:
        title_html = (
            '<h2 class="inventions-entry-title">'
            f'<span class="inventions-entry-num" aria-hidden="true">{esc(number)}</span>'
            f'<span class="inventions-entry-name">{esc(entry["title"])}</span></h2>'
        )
    else:
        title_html = (
            '<h2 class="inventions-entry-title">'
            f'<span class="inventions-entry-name">{esc(entry["title"])}</span></h2>'
        )
    parts = [
        f'<article class="inventions-entry" id="{esc(entry["slug"])}" '
        f'data-search="{esc(search_blob)}" '
        f'data-category="{esc(cat_num)}" '
        f'data-period="{esc(period)}">',
        title_html,
    ]
    visual_html = render_entry_visual(entry)
    if visual_html:
        parts.append(visual_html)
    if entry.get("meta"):
        parts.append(f'<p class="inventions-entry-meta">{esc(entry["meta"])}</p>')
    for key, label in SECTION_LABELS.items():
        if key == "MULTIMEDIA & FURTHER RESOURCES":
            continue
        lines = entry["sections"].get(key, [])
        if not lines:
            continue
        parts.append('<div class="inventions-entry-section">')
        parts.append(f"<h3>{esc(label)}</h3>")
        for line in lines:
            parts.append(f"<p>{esc(line)}</p>")
        parts.append("</div>")
    refs_html = render_entry_references(entry, ref_index, entry_refs_map)
    if refs_html:
        parts.append(refs_html)
    parts.append("</article>")
    return "\n".join(parts)


def _cat_number(title: str) -> str:
    m = re.match(r"(\d+)", title.strip())
    return m.group(1) if m else ""


def _cat_label(title: str) -> str:
    return re.sub(r"^\d+\.\s*", "", title).strip()


def render_toc(data: dict) -> str:
    items = ['<ul class="timeline-list" id="inventionsTocList">']
    for cat in data["categories"]:
        if not cat["entries"]:
            continue
        num = _cat_number(cat["title"])
        label = _cat_label(cat["title"])
        items.append(
            f'<li class="inventions-toc-cat-row" data-toc-cat="{esc(cat["slug"])}">'
            f'<span class="tl-date">§{esc(num)}</span>'
            f'<a href="#{esc(cat["slug"])}">{esc(label)}</a></li>'
        )
        for entry in cat["entries"]:
            number = entry.get("number", "")
            period = infer_period(entry.get("meta", ""), [], slug=entry.get("slug"))
            items.append(
                f'<li class="inventions-toc-entry" data-toc-entry="{esc(entry["slug"])}" '
                f'data-toc-cat="{esc(cat["slug"])}" '
                f'data-category="{esc(num)}" '
                f'data-period="{esc(period)}" '
                f'data-search="{esc((number + " " if number else "") + entry["title"])}">'
                f'<span class="tl-date">{esc(number) if number else "·"}</span>'
                f'<a href="#{esc(entry["slug"])}">{esc(entry["title"])}</a></li>'
            )
    items.append(
        '<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
        '<a href="#overview-by-category">Overview by category</a></li>'
    )
    if data["conclusion"]["title"]:
        items.append(
            f'<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
            f'<a href="#concluding-note">{esc(data["conclusion"]["title"])}</a></li>'
        )
    if data["references"]["title"]:
        items.append(
            f'<li class="inventions-toc-extra"><span class="tl-date">⊕</span>'
            f'<a href="#references">{esc(data["references"]["title"])}</a></li>'
        )
        for group in data["references"]["groups"]:
            label = re.sub(r"^[A-J]\.\s*", "", group["title"]).strip()
            items.append(
                f'<li class="inventions-toc-ref-row" data-search="{esc(group["title"])}">'
                f'<span class="tl-date">↳</span>'
                f'<a href="#{esc(group["slug"])}">{esc(label)}</a></li>'
            )
    items.append("</ul>")
    return "\n".join(items)


def render_references(data: dict) -> str:
    refs = data["references"]
    if not refs["title"] and not refs["groups"]:
        return ""

    parts = [
        '<section class="inventions-references" id="references">',
        f'<h2>{esc(refs["title"] or "References")}</h2>',
    ]
    if refs.get("intro"):
        parts.append(f'<p class="inventions-references-intro">{esc(refs["intro"])}</p>')

    for group in refs["groups"]:
        parts.append(
            f'<div class="inventions-references-group" id="{esc(group["slug"])}">'
            f'<h3>{esc(group["title"])}</h3>'
            '<ol class="inventions-bibliography">'
        )
        for item in group["items"]:
            parts.append(f"<li>{linkify_text(item)}</li>")
        parts.append("</ol></div>")

    parts.append("</section>")
    return "\n".join(parts)


def render_table(rows: list[list[str]], caption: str) -> str:
    if not rows:
        return ""
    head, body = rows[0], rows[1:]
    th = "".join(f"<th>{esc(c)}</th>" for c in head)
    trs = []
    for row in body:
        tds = "".join(f"<td>{esc(c)}</td>" for c in row)
        trs.append(f"<tr>{tds}</tr>")
    return f"""
<div class="inventions-table-wrap">
  <h2>{esc(caption)}</h2>
  <table class="inventions-data-table">
    <thead><tr>{th}</tr></thead>
    <tbody>{"".join(trs)}</tbody>
  </table>
</div>"""


def build_html(data: dict) -> str:
    nav = load_nav_html()
    ref_index = build_reference_index(data)
    entry_refs_map = load_entry_references()
    categories_html = []
    for cat in data["categories"]:
        if not cat["entries"]:
            continue
        num = _cat_number(cat["title"])
        entries_html = "\n".join(
            render_entry(e, ref_index, entry_refs_map, num) for e in cat["entries"]
        )
        categories_html.append(
            f'<section class="inventions-category" id="{esc(cat["slug"])}" '
            f'data-category="{esc(num)}">'
            f'<h2 class="inventions-category-head">{esc(cat["title"])}</h2>'
            f"{entries_html}</section>"
        )

    conclusion_html = ""
    if data["conclusion"]["paragraphs"]:
        ps = "".join(f"<p>{esc(p)}</p>" for p in data["conclusion"]["paragraphs"])
        conclusion_html = f"""
<section class="inventions-conclusion" id="concluding-note">
  <h2>{esc(data["conclusion"]["title"] or "Concluding note")}</h2>
  {ps}
</section>"""

    exec_summary = esc(data["executive_summary"] or data["tagline"])

    return f"""<!DOCTYPE html>
<html lang="en" data-kt-lang="en" data-kt-asset-root="../" data-kt-page-id="major-scientific-inventions" data-kt-nav-mount="1">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover"/>
{favicon_links("../")}
<title>Knowledge Treasury — Major Scientific Inventions</title>
<meta name="description" content="Sixty landmark inventions, discoveries, and innovations that shaped human history — an expanded reference with historical context, scientific significance, and curated multimedia resources."/>
<link rel="canonical" href="https://bilik-xezinesi.az/en/major_scientific_inventions.html"/>
<link rel="alternate" hreflang="az" href="https://bilik-xezinesi.az/az/major_scientific_inventions.html"/>
<link rel="alternate" hreflang="en" href="https://bilik-xezinesi.az/en/major_scientific_inventions.html"/>
<link rel="alternate" hreflang="x-default" href="https://bilik-xezinesi.az/az/major_scientific_inventions.html"/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link href="https://fonts.googleapis.com/css2?family=Inter:opsz,wght@14..32,400..900&amp;family=Playfair+Display:wght@700;800&amp;display=swap" rel="stylesheet"/>
<link href="../css/kt-common.css?v=70" rel="stylesheet"/>
<link href="../css/kt-perf.css?v=1" rel="stylesheet"/>
<link href="../css/kt-mobile.css?v=13" rel="stylesheet"/>
<link href="../css/kt-sticky-chrome.css?v=1" rel="stylesheet"/>
<link href="../css/kt-back-to-top.css?v=2" rel="stylesheet"/>
<link href="../css/kt-lang.css?v=13" rel="stylesheet"/>
<link href="../css/kt-nav-mega.css?v=69" rel="stylesheet"/>
<link href="../css/kt-hero-summary.css?v=13" rel="stylesheet"/>
<link href="../css/kt-content-hero.css?v=7" rel="stylesheet"/>
<link href="../css/kt-sidebar-widget.css?v=7" rel="stylesheet"/>
<link href="../css/kt-catalog-toolbar.css?v=15" rel="stylesheet"/>
<link href="../css/kt-inventions.css?v=25" rel="stylesheet"/>
<script src="../js/kt-mobile.js?v=6" defer></script>
<script src="../js/kt-perf.js?v=1" defer></script>
<script src="../js/kt-sticky-chrome.js?v=3" defer></script>
<script src="../js/kt-back-to-top.js?v=3" defer></script>
<script src="../js/kt-i18n.js?v=32" defer></script>
<script src="../js/kt-lang-position.js?v=7" defer></script>
<script src="../js/kt-design-tokens.js?v=2" defer></script>
<script src="../js/kt-nav.js?v=31" defer></script>
<script src="../js/kt-primary-nav.js?v=56" defer></script>
<script src="../js/kt-breadcrumbs.js?v=34" defer></script>
<script src="../js/kt-shell.js?v=13" defer></script>
<script src="../js/kt-page-subtitle.js?v=2" defer></script>
<script src="../js/kt-catalog-multi-filter.js?v=2" defer></script>
<script src="../js/kt-catalog-toolbar-mobile.js?v=8" defer></script>
<script src="../js/kt-historical-periods.js?v=1" defer></script>
<script src="../js/kt-sidebar-toc-groups.js?v=7" defer></script>
<script src="../js/kt-inventions.js?v=12" defer></script>
</head>
<body class="charter-page inventions-preview-page">
<a class="skip" href="#content">Skip to content</a>
{nav}
<header class="hero page-hero"><div class="hero-wrap shell"><section><h1 aria-describedby="page-hero-subtitle">Major Scientific Inventions</h1><p class="page-hero-subtitle" id="page-hero-subtitle" role="doc-subtitle">{esc(data["tagline"] or data["subtitle"] or "Discoveries and innovations in human history")}</p></section></div></header>
<main class="main" id="content">
<div class="toolbar catalog-toolbar" aria-label="Catalog filters">
<div class="catalog-toolbar__head">
<div class="search-wrap">
<svg fill="none" height="15" stroke="currentColor" stroke-width="2" viewBox="0 0 24 24" width="15" aria-hidden="true"><circle cx="11" cy="11" r="8"></circle><line x1="21" x2="16.65" y1="21" y2="16.65"></line></svg>
<input type="search" id="inventionsSearch" placeholder="Search inventions, periods, or key figures…" aria-label="Search inventions" autocomplete="off"/>
</div>
<button type="button" class="catalog-toolbar__toggle" aria-expanded="false" aria-controls="catalogFilterPanel" aria-label="Show filters"><svg class="catalog-toolbar__toggle-icon" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true"><path d="M3 5h18l-7 9v5l-4-2v-3L3 5z"/></svg><span class="catalog-toolbar__toggle-text">Filters</span><span class="catalog-toolbar__badge" hidden></span></button>
</div>
<div class="catalog-toolbar__panel" id="catalogFilterPanel">
<div class="catalog-toolbar__panel-inner">
<div class="filter-group">
<span class="filter-icon-label" aria-hidden="true"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M3 5h18l-7 9v5l-4-2v-3L3 5z"/></svg><span>Filters</span></span>
<div class="filter-group__primary">
<div class="sel-wrap">
<select id="filterCategory" aria-label="Category">{render_category_select(data)}</select>
<button class="sel-clear" data-for="filterCategory" title="Clear filter" type="button">×</button>
</div>
<div class="sel-wrap">
<select id="filterPeriod" aria-label="Historical period">{render_period_select()}</select>
<button class="sel-clear" data-for="filterPeriod" title="Clear filter" type="button">×</button>
</div>
</div>
<button class="btn-clear" id="clearFilters" type="button">Clear all</button>
</div>
</div>
</div>
</div>
<div class="charter-layout inventions-layout">
<aside class="charter-sidebar toc-card" aria-label="Article navigation">
<div class="sidebar-widget events-open" id="inventionsArticlesWidget">
<div class="widget-head">
<h2 class="widget-head__title"><span class="widget-head__icon" aria-hidden="true">📚</span> Articles</h2>
<button type="button" class="events-menu-toggle" aria-expanded="true" aria-controls="inventionsArticlesWidgetBody" aria-label="Toggle articles menu"><span></span><span></span><span></span></button>
</div>
<div class="widget-actions" role="group" aria-label="Articles section controls">
<button type="button" class="widget-action-btn widget-action-btn--toggle widget-action-btn--will-collapse" data-toc-action="toggle-categories" data-bulk-kind="categories" data-bulk-action="collapse" aria-pressed="true" aria-expanded="true" aria-label="Collapse all categories" title="Collapse All Categories"><span class="widget-action-btn__icon" aria-hidden="true"></span><span class="widget-action-btn__label">Collapse All Categories</span></button>
</div>
<div class="widget-body" id="inventionsArticlesWidgetBody">
{render_toc(data)}
</div>
</div>
</aside>
<div class="charter-stack inventions-stack">
{"".join(categories_html)}
<div id="overview-by-category">{render_table(data["overview_table"], "Overview by category")}</div>
{conclusion_html}
{render_references(data)}
</div>
</div>
</main>
<footer class="footer-pro">
<div class="footer-inner">
<div class="footer-brand"><h3>Knowledge Treasury</h3></div>
<div class="footer-grid">
<div class="footer-col"><div class="footer-title">Contact</div><div class="footer-item"><span aria-hidden="true">✉</span> <a href="mailto:info@bilik-xezinesi.az">info@bilik-xezinesi.az</a></div></div>
</div>
</div>
<div class="footer-bottom">© 2026 Knowledge Treasury — All rights reserved</div>
</footer>
</body>
</html>
"""


def verify_entry_references(data: dict, html_out: str, entry_refs_map: dict[str, list[str]]) -> None:
    missing: list[str] = []
    for entry in data["entries"]:
        slug = entry["slug"]
        if not entry_refs_map.get(slug):
            missing.append(f"{entry.get('number', '')} {slug} (no mapping)")
            continue
        block_id = f'id="{slug}"'
        if block_id not in html_out:
            missing.append(f"{slug} (article missing)")
            continue
        start = html_out.index(block_id)
        end = html_out.find('<article class="inventions-entry"', start + 1)
        if end == -1:
            end = len(html_out)
        block = html_out[start:end]
        if 'class="inventions-entry-references"' not in block:
            missing.append(f"{entry.get('number', '')} {slug} (no references block)")

    if missing:
        print(f"  Entry reference issues ({len(missing)}):")
        for line in missing:
            print(f"    {line}")
        raise SystemExit(1)
    print(f"  Entry reference blocks: {len(data['entries'])}")


def verify_entry_icons(data: dict, html_out: str) -> None:
    placeholders = len(re.findall(r'inventions-entry-icon-placeholder" role="img"', html_out))
    icon_images = len(re.findall(r'inventions-entry-icon"><img src="\.\./images/icons/', html_out))
    invention_images = len(
        re.findall(r'inventions-entry-icon"><img src="\.\./images/inventions/', html_out)
    )
    missing: list[str] = []
    mismatched: list[str] = []

    for entry in data["entries"]:
        slug = entry["slug"]
        number = entry.get("number", "")
        block_id = f'id="{slug}"'
        if block_id not in html_out:
            continue
        start = html_out.index(block_id)
        end = html_out.find('<article class="inventions-entry"', start + 1)
        if end == -1:
            end = len(html_out)
        block = html_out[start:end]

        if number.startswith(f"{ILLUSTRATION_SECTION}."):
            if "../images/inventions/" not in block:
                missing.append(f"{number} {slug} (section 6 illustration)")
            continue

        icon_src = entry.get("card_icon", "")
        if not icon_src or not (ROOT / icon_src.removeprefix("../")).is_file():
            missing.append(f"{number} {slug} (icon file)")
            continue
        if icon_src not in block:
            mismatched.append(f"{number} {slug}")

    print(f"  Icon images: {icon_images}; section {ILLUSTRATION_SECTION} illustrations: {invention_images}")
    print(f"  Placeholders remaining: {placeholders}")
    if missing:
        print(f"  Missing ({len(missing)}):")
        for line in missing:
            print(f"    {line}")
    if mismatched:
        print(f"  Mismatched ({len(mismatched)}):")
        for line in mismatched:
            print(f"    {line}")
    if placeholders or missing or mismatched:
        raise SystemExit(1)


def main() -> None:
    data = parse_docx()
    cards = load_card_data()
    attach_orphan_entries(data, cards)
    apply_link_fixes(data)
    attach_card_content(data, cards, load_card_overrides(), load_card_copy())
    entry_refs_map = load_entry_references()
    html_out = build_html(data)
    OUT_HTML.write_text(html_out, encoding="utf-8", newline="\n")
    az_out = html_out.replace('lang="en"', 'lang="az"', 1)
    az_out = az_out.replace("Main navigation", "Əsas naviqasiya", 1)
    OUT_HTML_AZ.write_text(az_out, encoding="utf-8", newline="\n")
    verify_entry_references(data, html_out, entry_refs_map)
    verify_entry_icons(data, html_out)
    print(f"Wrote {OUT_HTML.relative_to(ROOT)}")
    print(f"Wrote {OUT_HTML_AZ.relative_to(ROOT)}")
    print(f"  Categories: {len(data['categories'])}")
    print(f"  Entries: {len(data['entries'])}")


if __name__ == "__main__":
    main()
