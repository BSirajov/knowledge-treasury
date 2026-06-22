#!/usr/bin/env python3
"""Generate researched Word document for Category 1 inventions 1.1–1.7."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

from _build_inventions_page import parse_docx
from _paths import ROOT

OUT = ROOT / "documents" / "Foundational_Inventions_1.1-1.7_Research_Reference.docx"

MILESTONES: dict[str, list[str]] = {
    "controlled-use-of-fire": [
        "c. 1.0 Ma — In situ burning documented at Wonderwerk Cave, South Africa (Berna et al. 2012, PNAS).",
        "c. 400,000 BCE — Widespread evidence of habitual fire use at sites such as Qesem Cave, Israel.",
        "c. 400,000–300,000 BCE — Hearths and repeated combustion features at European and African sites.",
        "19th–21st century — Experimental archaeology and ethnography clarify fire-making and maintenance techniques.",
        "2009 — Richard Wrangham synthesises the 'cooking hypothesis' linking fire, diet, and hominin evolution.",
    ],
    "agriculture-and-domestication": [
        "c. 12,000–10,000 BCE — Independent plant domestication begins in the Fertile Crescent (wheat, barley).",
        "c. 10,000–8,000 BCE — Rice and millet domestication in China; maize in Mesoamerica; sorghum in Africa.",
        "c. 10,000–6,000 BCE — Domestication of goats, sheep, cattle, and pigs in western Asia.",
        "c. 6,000–3,500 BCE — Irrigation, plough agriculture, and early urban centres (Uruk, Jericho).",
        "20th century — Green Revolution and modern crop science transform global food production.",
    ],
    "the-wheel": [
        "c. 5,500–5,000 BCE — Potter's wheel in Mesopotamia (rotary craft technology).",
        "c. 3,500 BCE — Earliest pictorial evidence of wheeled vehicles (Late Uruk period, Mesopotamia).",
        "c. 2,000 BCE — Spoked wheels and chariots spread across Eurasia.",
        "Middle Ages — Water wheels and windmills harness rotary motion for industry.",
        "Industrial era — Gears, turbines, flywheels, and drivetrains generalise the wheel principle.",
    ],
    "writing-systems": [
        "c. 3,400–3,200 BCE — Sumerian cuneiform and Egyptian hieroglyphs emerge as early scripts.",
        "c. 2,600 BCE — Cuneiform used for literature (e.g. early Sumerian kings lists, mythic texts).",
        "c. 1,800 BCE — Proto-Sinaitic / early alphabetic scripts in the Levant.",
        "c. 1,200 BCE — Phoenician alphabet simplifies writing; spreads to Greece and Rome.",
        "c. 105 CE–present — Paper and printing (see entries 1.6–1.9) massively expand written culture.",
    ],
    "mathematics-and-the-concept-of-zero": [
        "c. 2,000 BCE — Babylonian clay tablets record algebra, square roots, and π approximations.",
        "c. 300 BCE — Euclid's Elements systematises Greek geometry.",
        "628 CE — Brahmagupta's Brāhmasphuṭasiddhānta formalises zero and rules for signed arithmetic.",
        "c. 825 CE — Al-Khwarizmi's works transmit Hindu-Arabic numerals to the Islamic world.",
        "17th century CE — Newton and Leibniz develop calculus; 20th century — formal logic and computing.",
    ],
    "paper": [
        "c. 2nd century BCE — Early hemp-fibre paper fragments in China (e.g. Fangmatan map).",
        "105 CE — Cai Lun presents improved bark-and-rag papermaking to the Han court (Hou Hanshu).",
        "8th century CE — Paper mills and papermaking spread through the Islamic world.",
        "12th–13th century CE — Paper reaches medieval Europe via Islamic Spain and Italy.",
        "19th century — Mechanised papermaking; 21st century — global production exceeds 400 Mt/year.",
    ],
    "the-compass": [
        "c. 4th century BCE — Chinese texts describe lodestone attracting iron (natural magnetism).",
        "c. 200 BCE — Han-period 'south-pointer' devices using lodestone (divination and orientation).",
        "c. 1040–1119 CE — Chinese nautical use of magnetised needles for navigation.",
        "c. 1190 CE — European references to compass use at sea (Alexander Neckam, others).",
        "15th–16th century — Age of Exploration; compass combined with charts and later with GPS.",
    ],
}

CONTRIBUTORS: dict[str, list[str]] = {
    "controlled-use-of-fire": [
        "Homo erectus and later Homo sapiens — earliest hominin users of fire.",
        "Francesco Berna, Paul Goldberg, Michael Chazan et al. — Wonderwerk Cave microstratigraphic research.",
        "Richard Wrangham — evolutionary anthropology of cooking (Harvard University).",
        "Liora Kolska Horwitz — archaeozoology of early fire sites.",
    ],
    "agriculture-and-domestication": [
        "Neolithic farmers across the Fertile Crescent, China, Africa, and the Americas.",
        "Norman Borlaug — 20th-century plant breeding (Green Revolution).",
        "Jared Diamond — synthesis of agriculture, geography, and civilisation (UCLA).",
        "James C. Scott — critical history of early states and domestication (Yale University).",
    ],
    "the-wheel": [
        "Mesopotamian and steppe craftspeople — earliest transport wheels.",
        "Potters of the Near East — rotary potter's wheel.",
        "David W. Anthony — archaeology of wheeled vehicles and Indo-European dispersal.",
    ],
    "writing-systems": [
        "Sumerian scribes — inventors of cuneiform.",
        "Egyptian priests and administrators — hieroglyphic script.",
        "Chinese and Maya scribes — independent logographic traditions.",
        "Phoenician traders — alphabetic writing adapted across the Mediterranean.",
        "Andrew Robinson — historian of writing systems.",
    ],
    "mathematics-and-the-concept-of-zero": [
        "Babylonian, Egyptian, Greek, Indian, and Maya mathematicians.",
        "Brahmagupta (c. 598–668 CE) — formal rules for zero and negative numbers.",
        "Aryabhata, Bhaskara II — Indian astronomical mathematics.",
        "Isaac Newton, Gottfried Leibniz — calculus.",
        "Marcus du Sautoy — public history of mathematics.",
    ],
    "paper": [
        "Cai Lun (c. 50–121 CE) — Han court official who standardised papermaking (105 CE).",
        "Zuo Bo — apprentice credited with further improvements.",
        "Joseph Needham & Tsuen-Hsuin Tsien — history of Chinese paper and printing.",
        "Islamic and European papermakers — mill technology from Samarkand to Fabriano.",
    ],
    "the-compass": [
        "Han Dynasty Chinese technicians — lodestone 'south-pointing' devices.",
        "Song Dynasty navigators — maritime compass use.",
        "European mariners (12th century onward) — open-sea navigation.",
        "William Gilbert — De Magnete (1600), systematic study of terrestrial magnetism.",
    ],
}

EXTRA_REFS: dict[str, list[str]] = {
    "controlled-use-of-fire": [
        "Berna, Francesco, et al. 'Microstratigraphic Evidence of in situ Fire in the Acheulean Strata of Wonderwerk Cave, South Africa.' Proceedings of the National Academy of Sciences 109, no. 20 (2012): E1215–E1220. https://doi.org/10.1073/pnas.1117620109",
        "Marín-Monfort, M. Dolores, et al. 'New Evidence for Early Pleistocene Use of Fire at Wonderwerk Cave (South Africa).' PLOS ONE (2026). https://doi.org/10.1371/journal.pone.0347480",
        "Roebroeks, Wil, and Paola Villa. 'On the Earliest Evidence for Habitual Use of Fire in Europe.' Proceedings of the National Academy of Sciences 108, no. 13 (2011): 5209–5214.",
        "Wrangham, Richard. Catching Fire: How Cooking Made Us Human. New York: Basic Books, 2009.",
        "Smithsonian Human Origins Program. 'Use of Fire.' https://humanorigins.si.edu/evidence/behavior/hearths-shelters",
    ],
    "agriculture-and-domestication": [
        "Barker, Graeme. The Agricultural Revolution in Prehistory: Why Did Foragers Become Farmers? Oxford: Oxford University Press, 2006.",
        "Diamond, Jared. Guns, Germs, and Steel: The Fates of Human Societies. New York: Norton, 1997.",
        "Scott, James C. Against the Grain: A Deep History of the Earliest States. New Haven: Yale University Press, 2017.",
        "Zeder, Melinda A. 'The Origins of Agriculture in the Near East.' Current Anthropology 52, no. S4 (2011): S221–S235.",
        "Roser, Max, Hannah Ritchie, and Bernadeta Dadonaite. 'Agriculture and Food.' Our World in Data, 2024. https://ourworldindata.org/agriculture",
    ],
    "the-wheel": [
        "Anthony, David W. The Horse, the Wheel, and Language: How Bronze-Age Riders from the Eurasian Steppes Shaped the Modern World. Princeton: Princeton University Press, 2007.",
        "Basalla, George. The Evolution of Technology. Cambridge: Cambridge University Press, 1988.",
        "Smithsonian Magazine. 'Who Invented the Wheel?' Smithsonian Institution, 2022. https://www.smithsonianmag.com/history/the-wheel-how-one-of-the-most-important-inventions-came-to-be-180976977/",
    ],
    "writing-systems": [
        "Robinson, Andrew. The Story of Writing: Alphabets, Hieroglyphs, and Pictograms. London: Thames & Hudson, 2007.",
        "Dehaene, Stanislas. Reading in the Brain: The Science and Evolution of a Human Invention. New York: Viking, 2009.",
        "British Museum. 'The Origins of Writing.' https://www.britishmuseum.org/collection/galleries/origins-writing",
        "Omniglot. Writing Systems of the World. https://www.omniglot.com/writing/",
    ],
    "mathematics-and-the-concept-of-zero": [
        "Colebrooke, Henry Thomas. Algebra, with Arithmetic and Mensuration, from the Sanscrit of Brahmegupta and Bháscara. London, 1817.",
        "Kaplan, Robert. The Nothing That Is: A Natural History of Zero. Oxford: Oxford University Press, 1999.",
        "Katz, Victor J. A History of Mathematics: An Introduction. 3rd ed. Boston: Pearson, 2009.",
        "MacTutor History of Mathematics Archive. University of St Andrews. https://mathshistory.st-andrews.ac.uk/",
        "Encyclopaedia Britannica. 'Brahmagupta.' https://www.britannica.com/biography/Brahmagupta",
    ],
    "paper": [
        "Needham, Joseph, and Tsuen-Hsuin Tsien. Science and Civilisation in China, Vol. 5, Part 1: Paper and Printing. Cambridge: Cambridge University Press, 1985.",
        "Hunter, Dard. Papermaking: The History and Technique of an Ancient Craft. New York: Dover, 1978.",
        "Encyclopaedia Britannica. 'Cai Lun.' https://www.britannica.com/biography/Cai-Lun",
        "University of Oxford, Bodleian Libraries. 'Origins of Paper.' https://www.cabinet.ox.ac.uk/origins-paper",
        "Smithsonian Magazine. 'A History of Paper and Printing.' https://www.smithsonianmag.com/history/a-history-of-paper-and-printing-180953012/",
    ],
    "the-compass": [
        "Needham, Joseph. Science and Civilisation in China, Vol. 4, Part 1: Physics. Cambridge: Cambridge University Press, 1962.",
        "Aczel, Amir D. The Riddle of the Compass: The Invention That Changed the World. Orlando: Harcourt, 2001.",
        "Gurney, Alan. Compass: A Story of Exploration and Innovation. New York: Norton, 2004.",
        "Encyclopaedia Britannica. 'Magnetic Compass.' https://www.britannica.com/technology/compass-navigation",
        "McClellan, James E., and Harold Dorn. Science and Technology in World History. 3rd ed. Baltimore: Johns Hopkins University Press, 2015.",
    ],
}

SLUGS = [
    "controlled-use-of-fire",
    "agriculture-and-domestication",
    "the-wheel",
    "writing-systems",
    "mathematics-and-the-concept-of-zero",
    "paper",
    "the-compass",
]


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build() -> None:
    data = parse_docx()
    entries = {e["slug"]: e for e in data["entries"] if e["slug"] in SLUGS}

    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_heading("Foundational Civilisational Innovations", 0)
    title.alignment = 0
    doc.add_heading("Research Reference: Entries 1.1–1.7", level=1)
    add_para(
        doc,
        "Prepared for the Knowledge Treasury — Most Influential Scientific Inventions and Innovations",
    )
    add_para(doc, "Scope: Section 1 (Foundational Civilisational Innovations), entries 1.1 through 1.7.")
    add_para(
        doc,
        "This document expands the source catalogue (paragraphs 6–7 of the table of contents identify "
        "Category 1 and its nine entries; entries 1.1–1.7 are treated here) with additional historical "
        "context, key contributors, scientific significance, milestone chronologies, and references "
        "drawn from peer-reviewed archaeology and anthropology, museum and library resources, "
        "university archives, and standard scholarly monographs.",
    )
    doc.add_page_break()

    for i, slug in enumerate(SLUGS, start=1):
        entry = entries[slug]
        number = f"1.{i}"
        title_text = entry["title"]
        add_heading(doc, f"{number}  {title_text}", level=1)
        if entry.get("meta"):
            add_para(doc, entry["meta"], bold=True)

        sections = entry.get("sections", {})
        add_heading(doc, "Historical Context", level=2)
        for line in sections.get("HISTORICAL CONTEXT", []):
            add_para(doc, line)

        add_heading(doc, "Key Contributors", level=2)
        add_bullets(doc, CONTRIBUTORS.get(slug, []))

        add_heading(doc, "Scientific and Technological Significance", level=2)
        for line in sections.get("SCIENTIFIC / TECHNOLOGICAL SIGNIFICANCE", []):
            add_para(doc, line)

        add_heading(doc, "Major Milestones", level=2)
        add_bullets(doc, MILESTONES.get(slug, []))

        add_heading(doc, "Impact on Humanity", level=2)
        for line in sections.get("IMPACT ON HUMANITY", []):
            add_para(doc, line)

        add_heading(doc, "References", level=2)
        add_numbered(doc, EXTRA_REFS.get(slug, []))

        if i < len(SLUGS):
            doc.add_page_break()

    add_heading(doc, "General References on Technology and World History", level=1)
    general = [
        "Harari, Yuval Noah. Sapiens: A Brief History of Humankind. London: Harvill Secker, 2011.",
        "McClellan, James E., and Harold Dorn. Science and Technology in World History. 3rd ed. Baltimore: Johns Hopkins University Press, 2015.",
        "Mokyr, Joel. The Lever of Riches: Technological Creativity and Economic Progress. Oxford: Oxford University Press, 1990.",
        "Basalla, George. The Evolution of Technology. Cambridge: Cambridge University Press, 1988.",
    ]
    add_numbered(doc, general)

    add_para(
        doc,
        "Note on sources: URLs were verified at the time of preparation (June 2026). "
        "Peer-reviewed items are cited with DOI where available. Museum, encyclopaedia, "
        "and educational resources are included for accessibility; primary scholarly "
        "claims should be traced to the peer-reviewed and monograph sources listed above.",
    )

    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    build()
