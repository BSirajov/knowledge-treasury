#!/usr/bin/env python3
"""
Insert editable summary tables after each infographic image in the inventions docx.
Original images are preserved unchanged.
"""
from __future__ import annotations

import json
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "documents" / "Most_Influential_Scientific_Inventions_and_Innovations (with references).docx"
HTML_DATA = ROOT / "helpers" / "_inventions_table_data.json"
OCR_DATA = ROOT / "helpers" / "_inventions_table_data_ocr.json"

# Import sibling helper when run as script
import sys

sys.path.insert(0, str(ROOT / "helpers"))
from _analyze_inventions_docx import has_drawing, iter_block_items  # noqa: E402

CLOUD_COMPUTING = {
    "name": "Cloud Computing",
    "summary": (
        "Amazon Web Services (2006) offered on-demand computing as a utility. "
        "AWS, Azure, and Google Cloud now generate over $250B revenue annually "
        "and host most internet services."
    ),
    "key_facts": [
        "Salesforce SaaS pioneer: 1999",
        "Amazon Web Services launched: 2006",
        "Google Cloud: 2008; Microsoft Azure: 2010",
        "AWS+Azure+GCP: >$250B revenue/year (2024)",
        "Democratised access to AI and computing",
    ],
}

# Manual corrections where HTML/OCR both degraded infographic text
SUMMARY_OVERRIDES: dict[str, str] = {
    "agricultureanddomestication": (
        "Farming emerged independently in at least seven regions. Crop cultivation and "
        "animal domestication enabled permanent settlements, population growth, and the first cities."
    ),
    "paper": (
        "Cai Lun made paper from plant fibres in 105 CE. Paper reached the Islamic world "
        "by the 8th century, and Europe by the 12th — making knowledge cheaper and more portable."
    ),
}

FACT_OVERRIDES: dict[str, list[str]] = {
    "agricultureanddomestication": [
        "At least 7 independent origins worldwide",
        "First crops: wheat, barley, millet, rice",
        "First animals: goats, sheep, cattle, pigs",
        "Enabled cities, states, writing, trade",
        "Also introduced zoonotic diseases",
    ],
}

META_RE = re.compile(
    r"^Period:\s*(.+?)\s*\|\s*Key figure\(s\):\s*(.+)$",
    re.IGNORECASE | re.DOTALL,
)


def normalize_title(title: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", title.lower())


def load_html_lookup() -> dict[str, dict]:
    entries = json.loads(HTML_DATA.read_text(encoding="utf-8"))
    return {normalize_title(e["name"]): e for e in entries}


def load_ocr_by_index() -> dict[int, dict]:
    if not OCR_DATA.exists():
        return {}
    entries = json.loads(OCR_DATA.read_text(encoding="utf-8"))
    return {e["image_index"]: e for e in entries}


def clean_summary(text: str) -> str:
    text = re.sub(r"\bKEY\s*", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace(" -", " —").replace("~1M", "~1 million")
    return text


def clean_ocr_facts(facts: list[str]) -> list[str]:
    cleaned = []
    for fact in facts:
        f = fact.strip()
        if not f or f.upper() in {"ACTS", "KEY FACTS", "KEY", "FACTS"}:
            continue
        # repair common OCR truncations at line start
        repairs = {
            "est evidence": "Oldest evidence",
            "sites:": "Key sites:",
            "king fuelled": "Cooking fuelled",
            "bled settlement": "Enabled settlement",
            "cursor to": "Precursor to",
            "east 7": "At least 7",
            "t crops": "First crops",
            "t animals": "First animals",
            "bled cities": "Enabled cities",
            "t use:": "First use:",
            "nsport wheels": "Transport wheels",
        }
        for src, dst in repairs.items():
            if f.startswith(src):
                f = dst + f[len(src) :]
                break
        cleaned.append(f)
    return cleaned


def summary_from_raw_left(left: str) -> str:
    text = re.sub(r"Key figure\(s\):.*?(?:\n|$)", "", left, flags=re.I | re.S)
    parts = []
    for ln in text.splitlines():
        ln = ln.strip()
        if not ln or ln.upper() in {"KEY", "KEY FACTS", "FACTS", "ACTS"}:
            continue
        if len(ln) <= 2:
            continue
        parts.append(ln)
    return clean_summary(" ".join(parts))


def pick_summary(html_val: str, ocr_val: str, *, title_key: str = "", raw_left: str = "") -> str:
    if title_key in SUMMARY_OVERRIDES:
        return SUMMARY_OVERRIDES[title_key]
    raw_summary = summary_from_raw_left(raw_left) if raw_left else ""
    html_val = clean_summary(html_val or "")
    ocr_val = clean_summary(ocr_val or "")
    candidates = [c for c in (html_val, ocr_val, raw_summary) if c]
    if not candidates:
        return ""
    # prefer HTML unless it looks corrupted
    if html_val and not (
        re.search(r"[a-z]{3,}[A-Z]", html_val)
        or re.match(r"^[a-z]{2,4}\s", html_val)
        or "ivationand" in html_val
        or len(html_val) < 40
    ):
        return html_val
    # otherwise longest clean candidate
    return max(candidates, key=len)


def pick_facts(html_facts: list[str], ocr_facts: list[str], *, title_key: str = "") -> list[str]:
    if title_key in FACT_OVERRIDES:
        return FACT_OVERRIDES[title_key]
    html_facts = [f.strip().replace("barley.millet", "barley, millet") for f in html_facts if f.strip()]
    html_facts = [f.replace("writing. trade", "writing, trade") for f in html_facts]
    ocr_facts = clean_ocr_facts(ocr_facts)
    if len(html_facts) >= 4 and all(len(f) > 8 for f in html_facts[:3]):
        return html_facts
    return ocr_facts or html_facts


def parse_meta_line(text: str) -> tuple[str, str]:
    m = META_RE.match(text.strip())
    if not m:
        return "", ""
    return m.group(1).strip(), m.group(2).strip()


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def add_bullet_paragraph(cell, facts: list[str]) -> None:
    cell.text = ""
    for i, fact in enumerate(facts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.text = ""
        run = p.add_run(f"• {fact}")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.08)


def build_summary_table(doc: Document, entry: dict) -> Table:
    table = doc.add_table(rows=5, cols=2)
    table.style = "Normal Table"

    labels = ["Invention", "Period", "Key figure(s)", "Summary", "Key facts"]
    values = [
        entry["name"],
        entry["period"],
        entry["key_figures"],
        entry["summary"],
        "",
    ]

    for row_idx, (label, value) in enumerate(zip(labels, values)):
        label_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]
        set_cell_text(label_cell, label, bold=True)
        set_cell_shading(label_cell, "D9EAF7")
        if row_idx == 4:
            add_bullet_paragraph(value_cell, entry["key_facts"])
        else:
            set_cell_text(value_cell, value)

    for row in table.rows:
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.05)

    return table


def insert_element_after(paragraph: Paragraph, element) -> None:
    paragraph._element.addnext(element)


def insert_spacer_paragraph_after(paragraph: Paragraph) -> None:
    spacer = OxmlElement("w:p")
    paragraph._element.addnext(spacer)


def collect_articles(doc: Document) -> list[dict]:
    blocks = list(iter_block_items(doc))
    html_lookup = load_html_lookup()
    ocr_lookup = load_ocr_by_index()
    articles: list[dict] = []
    pending_title: str | None = None
    image_counter = 0

    for i, block in enumerate(blocks):
        if not isinstance(block, Paragraph):
            continue
        style = block.style.name if block.style else ""
        text = block.text.strip()
        if style == "Heading 2" and text:
            pending_title = text
        elif pending_title and has_drawing(block):
            image_counter += 1
            title = pending_title
            pending_title = None

            period, key_figures = "", ""
            for j in range(i + 1, min(i + 5, len(blocks))):
                nb = blocks[j]
                if isinstance(nb, Paragraph) and nb.text.strip():
                    period, key_figures = parse_meta_line(nb.text.strip())
                    if period:
                        break

            title_key = normalize_title(title)
            html = html_lookup.get(title_key, {})
            ocr = ocr_lookup.get(image_counter, {})

            if title_key == normalize_title("Cloud Computing"):
                summary = CLOUD_COMPUTING["summary"]
                key_facts = CLOUD_COMPUTING["key_facts"]
            else:
                summary = pick_summary(
                    html.get("summary", ""),
                    ocr.get("summary", ""),
                    title_key=title_key,
                    raw_left=ocr.get("raw_left", ""),
                )
                key_facts = pick_facts(
                    html.get("key_facts", []),
                    ocr.get("key_facts", []),
                    title_key=title_key,
                )

            if not period:
                period = html.get("period") or ocr.get("period", "")
            if not key_figures:
                key_figures = html.get("key_figures") or ocr.get("key_figures", "")

            articles.append(
                {
                    "title": title,
                    "image_index": i,
                    "entry": {
                        "name": title,
                        "period": period,
                        "key_figures": key_figures,
                        "summary": summary,
                        "key_facts": key_facts,
                    },
                }
            )

    return articles


def article_has_table_after(blocks: list, image_index: int) -> bool:
    for j in range(image_index + 1, min(image_index + 4, len(blocks))):
        if isinstance(blocks[j], Table):
            return True
    return False


def update_existing_table(table: Table, entry: dict) -> None:
    values = [
        entry["name"],
        entry["period"],
        entry["key_figures"],
        entry["summary"],
        "",
    ]
    labels = ["Invention", "Period", "Key figure(s)", "Summary", "Key facts"]
    for row_idx, value in enumerate(values):
        if row_idx == 4:
            add_bullet_paragraph(table.rows[row_idx].cells[1], entry["key_facts"])
        else:
            set_cell_text(table.rows[row_idx].cells[1], value)


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Missing document: {DOCX}")

    doc = Document(str(DOCX))
    blocks = list(iter_block_items(doc))
    articles = collect_articles(doc)

    # Update mode: refresh table cell text if tables already inserted
    updated = 0
    for item in articles:
        idx = item["image_index"]
        image_para = blocks[idx]
        for j in range(idx + 1, min(idx + 4, len(blocks))):
            nb = blocks[j]
            if isinstance(nb, Table) and nb.rows[0].cells[0].text == "Invention":
                update_existing_table(nb, item["entry"])
                updated += 1
                break

    if updated:
        doc.save(str(DOCX))
        print(f"Updated {updated} existing tables in {DOCX.name}")
        return

    backup = DOCX.with_suffix(f".backup-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx")
    shutil.copy2(DOCX, backup)
    print(f"Backup: {backup}")

    print(f"Prepared {len(articles)} article tables")

    for item in reversed(articles):
        image_para = blocks[item["image_index"]]
        table = build_summary_table(doc, item["entry"])
        tbl_el = table._tbl
        doc.element.body.remove(tbl_el)
        insert_spacer_paragraph_after(image_para)
        spacer_el = image_para._element.getnext()
        spacer_el.addnext(tbl_el)

    doc.save(str(DOCX))
    print(f"Saved {DOCX.name} with {len(articles)} editable tables (images preserved)")


if __name__ == "__main__":
    main()
