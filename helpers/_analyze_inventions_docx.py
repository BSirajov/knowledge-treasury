#!/usr/bin/env python3
"""Analyze structure of inventions docx: paragraphs, tables, inline images."""
from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "documents" / "Most_Influential_Scientific_Inventions_and_Innovations (with references).docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


def iter_block_items(parent):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    if isinstance(parent, DocumentType):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_style(p: Paragraph) -> str:
    return p.style.name if p.style else ""


def has_drawing(p: Paragraph) -> bool:
    return bool(p._element.findall(".//w:drawing", NS))


def drawing_extent_emu(p: Paragraph) -> tuple[int, int] | None:
    ext = p._element.find(".//wp:extent", NS)
    if ext is None:
        return None
    cx = int(ext.get("cx", 0))
    cy = int(ext.get("cy", 0))
    return cx, cy


def emu_to_in(cx: int, cy: int) -> tuple[float, float]:
    return cx / 914400, cy / 914400


def image_rids_in_paragraph(p: Paragraph) -> list[str]:
    rids = []
    for blip in p._element.findall(".//a:blip", NS):
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def main() -> None:
    doc = Document(str(DOCX))
    print(f"Document: {DOCX.name}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables (existing): {len(doc.tables)}")
    print()

    # media inventory
    with zipfile.ZipFile(DOCX) as zf:
        media = sorted(n for n in zf.namelist() if n.startswith("word/media/"))
        print(f"Media files: {len(media)}")
        for m in media[:5]:
            print(f"  {m}")
        if len(media) > 5:
            print(f"  ... +{len(media)-5} more")
    print()

    heading_re = re.compile(r"^\d+\.\d+\s")
    image_paras = []
    headings = []

    for i, block in enumerate(iter_block_items(doc)):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            style = paragraph_style(block)
            if heading_re.match(text) or style.startswith("Heading"):
                headings.append((i, style, text[:80]))
            if has_drawing(block):
                ext = drawing_extent_emu(block)
                w_in = h_in = 0.0
                if ext:
                    w_in, h_in = emu_to_in(*ext)
                rids = image_rids_in_paragraph(block)
                ctx_before = ""
                ctx_after = ""
                image_paras.append(
                    {
                        "idx": i,
                        "style": style,
                        "text": text[:60],
                        "w_in": round(w_in, 2),
                        "h_in": round(h_in, 2),
                        "rids": rids,
                    }
                )

    print(f"Headings found: {len(headings)}")
    for h in headings[:15]:
        print(f"  [{h[0]}] {h[1]}: {h[2]}")
    if len(headings) > 15:
        print(f"  ... +{len(headings)-15} more")
    print()

    print(f"Paragraphs with images: {len(image_paras)}")
    # classify by size - tables likely wider
    for p in image_paras[:30]:
        print(
            f"  [{p['idx']}] {p['w_in']}x{p['h_in']} in | style={p['style']} | text={p['text']!r}"
        )
    if len(image_paras) > 30:
        print(f"  ... +{len(image_paras)-30} more")

    # size histogram
    wide = [p for p in image_paras if p["w_in"] >= 4.0]
    tall = [p for p in image_paras if p["h_in"] >= 2.0 and p["w_in"] >= 3.0]
    small = [p for p in image_paras if p["w_in"] < 2.0]
    print()
    print(f"Wide images (w>=4in): {len(wide)}")
    print(f"Large images (w>=3, h>=2): {len(tall)}")
    print(f"Small images (w<2in): {len(small)}")


if __name__ == "__main__":
    main()
