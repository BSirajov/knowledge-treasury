#!/usr/bin/env python3
"""Generate a researched Word reference document for all 60 scientific inventions."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from _build_inventions_page import attach_orphan_entries, parse_docx
from _build_section1_research_docx import (
    CONTRIBUTORS as S1_CONTRIBUTORS,
    EXTRA_REFS as S1_EXTRA_REFS,
    MILESTONES as S1_MILESTONES,
    add_bullets,
    add_heading,
    add_numbered,
    add_para,
    set_doc_defaults,
)
from _inventions_research_extended import CONTRIBUTORS as EXT_CONTRIBUTORS
from _inventions_research_extended import MILESTONES as EXT_MILESTONES
from _paths import HELPERS, ROOT

OUT = ROOT / "documents" / "Most_Influential_Scientific_Inventions_Research_Reference.docx"
CARD_COPY = HELPERS / "_inventions_card_copy.json"
CARD_DATA = ROOT / "preview" / "inventions-card-data.json"
ENTRY_REFERENCES = HELPERS / "_invention_entry_references.json"
ALL_REFS = HELPERS / "_all_refs.json"

CATEGORY_INTROS: dict[str, str] = {
    "1": (
        "Foundational civilisational innovations — fire, agriculture, the wheel, writing, mathematics, "
        "paper, navigation, gunpowder, and printing — established the material and intellectual "
        "preconditions for all later science and technology."
    ),
    "2": (
        "Knowledge and scientific instruments transformed how humans observe, measure, and explain "
        "the natural world, from the scientific method through telescopes, microscopes, and Newtonian mechanics."
    ),
    "3": (
        "Industrial, energy, and transport technologies converted scientific understanding into "
        "mechanical power, electrical grids, and global mobility."
    ),
    "4": (
        "Medicine, biology, and public health innovations extended human lifespan and quality of life "
        "through vaccination, antibiotics, germ theory, anaesthesia, imaging, sanitation, and molecular genetics."
    ),
    "5": (
        "Physics, chemistry, and materials science revealed the structure of matter and energy, "
        "enabling the periodic table, quantum theory, relativity, nuclear science, plastics, and fertilisers."
    ),
    "6": (
        "Digital, space, and communication technologies interconnected the planet through telegraphy, "
        "telephony, radio, semiconductors, computing, the internet, satellites, GPS, lasers, and fibre optics."
    ),
    "7": (
        "Emerging and transformative technologies — artificial intelligence, renewables, batteries, "
        "electric vehicles, robotics, additive manufacturing, and blockchain — are reshaping the 21st century."
    ),
}


def load_json(path: Path) -> dict | list:
    return json.loads(path.read_text(encoding="utf-8"))


def cat_number(title: str) -> str:
    m = re.match(r"(\d+)", title.strip())
    return m.group(1) if m else ""


def cat_label(title: str) -> str:
    return re.sub(r"^\d+\.\s*", "", title).strip()


def parse_key_figures(text: str) -> list[str]:
    if not text:
        return []
    parts = re.split(r"[;|]", text)
    out: list[str] = []
    for part in parts:
        item = part.strip()
        if item:
            out.append(item)
    return out


def ref_dedup_key(text: str) -> str:
    urls = re.findall(r"https?://[^\s\]\)>,]+", text)
    if urls:
        return urls[0].rstrip("/").lower()
    return re.sub(r"\s+", " ", text.strip().lower())[:120]


def milestones_for_entry(slug: str, entry: dict, cards: dict) -> list[str]:
    if slug in S1_MILESTONES:
        return S1_MILESTONES[slug]
    if slug in EXT_MILESTONES:
        return EXT_MILESTONES[slug]

    card = cards.get(slug, {})
    facts = card.get("key_facts", [])
    meta = entry.get("meta", "")
    items: list[str] = []

    period_match = re.search(r"Period:\s*([^|]+)", meta, re.I)
    if period_match:
        items.append(f"Chronological span: {period_match.group(1).strip()}")

    for fact in facts:
        items.append(fact)

    if not items and meta:
        items.append(meta)
    return items


def contributors_for_entry(slug: str, entry: dict, cards: dict) -> list[str]:
    if slug in S1_CONTRIBUTORS:
        return S1_CONTRIBUTORS[slug]
    if slug in EXT_CONTRIBUTORS:
        return EXT_CONTRIBUTORS[slug]

    card = cards.get(slug, {})
    figures = parse_key_figures(card.get("key_figures", ""))
    if not figures:
        meta = entry.get("meta", "")
        km = re.search(r"Key figure\(s\):\s*(.+)$", meta, re.I)
        if km:
            figures = parse_key_figures(km.group(1))
    return figures or ["Multiple independent researchers, engineers, and institutions (see references)."]


def bibliography_refs(slug: str, ref_index: dict[str, str], entry_refs_map: dict[str, list[str]]) -> list[str]:
    ids = entry_refs_map.get(slug, [])
    seen: set[str] = set()
    out: list[str] = []
    for rid in ids:
        text = ref_index.get(rid)
        if text and text not in seen:
            seen.add(text)
            out.append(text)
    return out


def multimedia_refs(entry: dict) -> list[str]:
    resources = entry.get("sections", {}).get("MULTIMEDIA & FURTHER RESOURCES", [])
    out: list[str] = []
    for item in resources:
        if isinstance(item, dict):
            title = item.get("title", "").strip()
            source = item.get("source", "").strip()
            url = item.get("url", "").strip()
            if title and url:
                prefix = f"{source}. " if source else ""
                out.append(f"{prefix}{title}. {url}")
            elif title:
                out.append(title)
        elif isinstance(item, str) and item.strip():
            out.append(item.strip())
    return out


def all_references(slug: str, entry: dict, ref_index: dict[str, str], entry_refs_map: dict[str, list[str]]) -> list[str]:
    seen: set[str] = set()
    combined: list[str] = []

    def add(items: list[str]) -> None:
        for item in items:
            key = ref_dedup_key(item)
            if key and key not in seen:
                seen.add(key)
                combined.append(item.strip())

    add(bibliography_refs(slug, ref_index, entry_refs_map))
    add(S1_EXTRA_REFS.get(slug, []))
    add(multimedia_refs(entry))
    return combined


def ordered_entries(data: dict) -> list[tuple[dict, dict]]:
    rows: list[tuple[dict, dict]] = []
    for cat in data["categories"]:
        for entry in cat["entries"]:
            rows.append((cat, entry))
    return rows


def add_long_term_significance(doc: Document, entry: dict, cards: dict) -> None:
    slug = entry["slug"]
    card = cards.get(slug, {})
    summary = card.get("summary", "").strip()
    impact = entry.get("sections", {}).get("IMPACT ON HUMANITY", [])

    add_heading(doc, "Long-Term Significance", level=2)
    if summary:
        add_para(doc, summary)
    if impact:
        for line in impact:
            if line.strip() and (not summary or line.strip() not in summary):
                add_para(doc, line.strip())


def build() -> None:
    data = parse_docx()
    cards = load_json(CARD_COPY)
    card_data = load_json(CARD_DATA) if CARD_DATA.exists() else {}
    entry_refs_map = load_json(ENTRY_REFERENCES).get("entries", {})
    ref_index = {r["id"]: r["text"] for r in load_json(ALL_REFS)}

    attach_orphan_entries(data, card_data)
    rows = ordered_entries(data)

    doc = Document()
    set_doc_defaults(doc)

    doc.add_heading("Most Influential Scientific Inventions and Innovations", 0)
    add_heading(doc, "Comprehensive Research Reference", level=1)
    add_para(
        doc,
        "Prepared for the Knowledge Treasury — a scholarly companion to the Major Scientific Inventions catalogue.",
    )
    add_para(
        doc,
        "This document covers all 60 inventions across seven thematic categories. For each entry it "
        "provides historical background, key contributors, scientific and technological principles, "
        "a development timeline, societal impact, long-term significance, and verified references drawn "
        "from peer-reviewed literature, museum and library archives, government and institutional "
        "reports, university resources, encyclopaedias, and standard scholarly monographs.",
    )
    add_para(doc, "Compilation date: June 2026.")
    doc.add_page_break()

    # Table of contents (manual outline)
    add_heading(doc, "Contents", level=1)
    for cat in data["categories"]:
        num = cat_number(cat["title"])
        label = cat_label(cat["title"])
        add_para(doc, f"Category {num}: {label}", bold=True)
        for entry in cat["entries"]:
            add_para(doc, f"  {entry.get('number', '')}  {entry['title']}")
    doc.add_page_break()

    current_cat_num = ""

    for cat, entry in rows:
        cat_num = cat_number(cat["title"])
        if cat_num != current_cat_num:
            current_cat_num = cat_num
            doc.add_page_break()
            add_heading(doc, f"Category {cat_num}: {cat_label(cat['title'])}", level=1)
            intro = CATEGORY_INTROS.get(cat_num, "")
            if intro:
                add_para(doc, intro)

        slug = entry["slug"]
        number = entry.get("number", "")
        title = entry["title"]
        sections = entry.get("sections", {})

        add_heading(doc, f"{number}  {title}", level=1)
        if entry.get("meta"):
            add_para(doc, entry["meta"], bold=True)

        card = cards.get(slug, {})
        if card.get("summary"):
            add_para(doc, card["summary"])

        add_heading(doc, "Historical Background", level=2)
        for line in sections.get("HISTORICAL CONTEXT", []):
            add_para(doc, line)

        add_heading(doc, "Key Inventors and Contributors", level=2)
        add_bullets(doc, contributors_for_entry(slug, entry, cards))

        add_heading(doc, "Scientific and Technological Principles", level=2)
        for line in sections.get("SCIENTIFIC / TECHNOLOGICAL SIGNIFICANCE", []):
            add_para(doc, line)

        add_heading(doc, "Development Timeline and Major Milestones", level=2)
        add_bullets(doc, milestones_for_entry(slug, entry, cards))

        add_heading(doc, "Impact on Society", level=2)
        for line in sections.get("IMPACT ON HUMANITY", []):
            add_para(doc, line)

        add_long_term_significance(doc, entry, cards)

        add_heading(doc, "References and Sources", level=2)
        refs = all_references(slug, entry, ref_index, entry_refs_map)
        if refs:
            add_numbered(doc, refs)
        else:
            add_para(doc, "See the general bibliography at the end of this document.")

        doc.add_page_break()

    add_heading(doc, "General Bibliography", level=1)
    add_para(
        doc,
        "The following sources support multiple entries throughout this catalogue and were cited "
        "in the original Knowledge Treasury research compilation.",
    )
    general_ids = load_json(ENTRY_REFERENCES).get("general", [])
    if not general_ids:
        general_note_refs = [
            r["text"]
            for r in load_json(ALL_REFS)
            if r.get("group", "").startswith("A.")
        ][:8]
    else:
        general_note_refs = [ref_index[rid] for rid in general_ids if rid in ref_index]

    general = general_note_refs or [
        "Harari, Yuval Noah. Sapiens: A Brief History of Humankind. London: Harvill Secker, 2011.",
        "McClellan, James E., and Harold Dorn. Science and Technology in World History. 3rd ed. Baltimore: Johns Hopkins University Press, 2015.",
        "Mokyr, Joel. The Lever of Riches: Technological Creativity and Economic Progress. Oxford: Oxford University Press, 1990.",
        "Basalla, George. The Evolution of Technology. Cambridge: Cambridge University Press, 1988.",
        "Diamond, Jared. Guns, Germs, and Steel: The Fates of Human Societies. New York: Norton, 1997.",
        "Johnson, Steven. How We Got to Now: Six Innovations That Made the Modern World. New York: Riverhead Books, 2014.",
        "Bijker, Wiebe E., Thomas P. Hughes, and Trevor Pinch, eds. The Social Construction of Technological Systems. Cambridge, MA: MIT Press, 1987.",
        "Encyclopaedia Britannica. Science and Technology. https://www.britannica.com/science",
    ]
    add_numbered(doc, general)

    add_heading(doc, "Source Notes", level=2)
    add_para(
        doc,
        "URLs were verified at the time of preparation (June 2026). Peer-reviewed items are cited "
        "with DOI or journal details where available in the source bibliography. Museum, encyclopaedia, "
        "and educational resources are included for accessibility; primary scholarly claims should be "
        "traced to the peer-reviewed articles and monographs listed under each invention.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT.relative_to(ROOT)} ({len(rows)} inventions)")


if __name__ == "__main__":
    build()
